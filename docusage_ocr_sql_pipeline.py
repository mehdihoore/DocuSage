import os
import re
import ssl
import time
import json
import fitz
import httpx
import hashlib
from google import genai
from docx import Document
from mistralai.client import Mistral
from google.colab import userdata

# ==========================================
# CONFIGURATION
# 1 MISTRAL_API_KEY- 2- Misteralphal==========================================
MISTRAL_API_KEY = userdata.get('MistralHo')

GEMINI_API_KEYS = [
    userdata.get('GOOGLE_API_KEY1'),
    userdata.get('GOOGLE_API_KEY_1'),
    userdata.get('GOOGLE_API_KEY_2'),
    userdata.get('GOOGLE_API_KEY_3'),
    userdata.get('GOOGLE_API_KEY_4'),
    userdata.get('GOOGLE_API_KEY_BHR'),
    userdata.get('GOOGLE_API_KEY_BHR1'),
    userdata.get('GOOGLE_SUMMARY_API_KEY'),
    userdata.get('GEMINI_API_KEY'),
]
GEMINI_API_KEYS = [k for k in GEMINI_API_KEYS if k]

if not GEMINI_API_KEYS:
    raise ValueError("No Gemini API keys found in Colab secrets!")

print(f"Loaded {len(GEMINI_API_KEYS)} Gemini API key(s).")

GEMINI_MODEL_FALLBACK_CHAIN = [
    "gemini-2.0-flash",
    "gemma-4-31b-it",
    "gemma-4-26b-a4b-it",
    "gemini-1.5-flash-8b",
]

_gemini_key_index = 0

INPUT_BOOKS_DIR = r"/content/drive/MyDrive/Philosophy/Unsorted_Books"
OUTPUT_DIR      = r"/content/drive/MyDrive/Philosophy/Processed"

MAX_FILE_SIZE_MB     = 95.0
MAX_PAGES_PER_MISTRAL = 1000
PAGES_PER_CHUNK      = 400
SLEEP_BETWEEN_CHUNKS = 5
SLEEP_BETWEEN_BOOKS  = 10

SQL_DUMP_FILE_PATH   = os.path.join(OUTPUT_DIR, "docusage_insert_dump.sql")
FINGERPRINT_DB_PATH  = os.path.join(OUTPUT_DIR, ".processed_fingerprints.json")

# ==========================================
# GEMINI CLIENT HELPERS
# ==========================================
def get_active_gemini_client():
    return genai.Client(api_key=GEMINI_API_KEYS[_gemini_key_index])

def rotate_gemini_key(error_msg: str = ""):
    global _gemini_key_index
    _gemini_key_index += 1
    if _gemini_key_index >= len(GEMINI_API_KEYS):
        raise RuntimeError(
            f"All {len(GEMINI_API_KEYS)} Gemini API keys exhausted. Last error: {error_msg}"
        )
    print(f"  [Gemini] Rotated to key #{_gemini_key_index + 1} of {len(GEMINI_API_KEYS)}.")

def is_quota_error(error_str: str) -> bool:
    return any(kw in error_str.lower() for kw in [
        "quota", "exhausted", "resource_exhausted",
        "429", "403", "forbidden", "permission", "api key", "invalid", "access"
    ])

def is_daily_quota_error(error_str: str) -> bool:
    return ("perday" in error_str.lower() or
            "per_day" in error_str.lower() or
            "GenerateRequestsPerDay" in error_str)

def parse_gemini_retry_after(error_str: str) -> float:
    match = re.search(r"retryDelay.*?'(\d+)s'", error_str)
    if match:
        return float(match.group(1))
    match = re.search(r"retry in (\d+(?:\.\d+)?)s", error_str, re.IGNORECASE)
    if match:
        return float(match.group(1))
    return 0.0

def try_gemini_model(model_name: str, prompt: str) -> dict:
    """Tries all API keys for a single model. Returns dict or raises RuntimeError."""
    global _gemini_key_index
    keys_tried = 0

    while keys_tried < len(GEMINI_API_KEYS):
        current_key_num = _gemini_key_index + 1
        try:
            gemini_client = get_active_gemini_client()
            print(f"  [{model_name}] Trying key #{current_key_num}/{len(GEMINI_API_KEYS)}...")
            response = gemini_client.models.generate_content(
                model=model_name,
                contents=prompt,
                config={"response_mime_type": "application/json"}
            )
            result = json.loads(response.text.strip())
            print(f"  [{model_name}] ✓ Success with key #{current_key_num}.")
            return result

        except Exception as e:
            error_str = str(e)
            if not is_quota_error(error_str):
                print(f"  [{model_name}] Non-quota error: {error_str[:120]}")
                raise RuntimeError(f"Non-quota failure on {model_name}: {error_str}")

            if is_daily_quota_error(error_str):
                print(f"  [{model_name}] Key #{current_key_num} daily quota=0. Rotating...")
            else:
                retry_after = parse_gemini_retry_after(error_str)
                if 0 < retry_after <= 15:
                    print(f"  [{model_name}] Rate limited. Waiting {retry_after:.0f}s...")
                    time.sleep(retry_after)
                else:
                    print(f"  [{model_name}] Rate limited ({retry_after:.0f}s — too long). Rotating...")

            try:
                rotate_gemini_key(error_str)
                keys_tried += 1
            except RuntimeError:
                raise RuntimeError(f"All {len(GEMINI_API_KEYS)} keys exhausted for {model_name}.")

    raise RuntimeError(f"All keys failed for {model_name}.")

def call_llm(prompt: str, context: str = "") -> dict:
    """
    Universal LLM caller used by both extract_metadata() and check_duplicate_via_llm().
    Tries full Gemini fallback chain, then Mistral. Returns parsed JSON dict.
    """
    global _gemini_key_index

    for model_name in GEMINI_MODEL_FALLBACK_CHAIN:
        _gemini_key_index = 0
        try:
            return try_gemini_model(model_name, prompt)
        except RuntimeError as e:
            print(f"  [LLM] {model_name} exhausted: {e}. Trying next...")
            continue

    # Mistral fallback
    print(f"  [LLM] All Gemini models exhausted{' for ' + context if context else ''}. Trying Mistral...")
    try:
        response = client.chat.complete(
            model="mistral-small-latest",
            messages=[{"role": "user", "content": prompt}]
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.IGNORECASE)
        raw = re.sub(r'\s*```$', '', raw)
        return json.loads(raw)
    except Exception as e:
        raise RuntimeError(f"All LLMs failed: {e}")

# ==========================================
# MISTRAL CLIENT
# ==========================================
ssl_context = ssl.create_default_context()
ssl_context.minimum_version = ssl.TLSVersion.TLSv1_2
ssl_context.maximum_version = ssl.TLSVersion.TLSv1_2

custom_http_client = httpx.Client(
    verify=ssl_context,
    trust_env=True,
    timeout=None
)

client = Mistral(
    api_key=MISTRAL_API_KEY,
    client=custom_http_client
)

# ==========================================
# FINGERPRINT DATABASE
# ==========================================
def load_fingerprint_db() -> dict:
    if not os.path.exists(FINGERPRINT_DB_PATH):
        return {}
    try:
        with open(FINGERPRINT_DB_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def save_fingerprint_db(db: dict):
    with open(FINGERPRINT_DB_PATH, "w", encoding="utf-8") as f:
        json.dump(db, f, ensure_ascii=False, indent=2)

def get_file_hash(pdf_path: str) -> str:
    hasher = hashlib.sha256()
    with open(pdf_path, "rb") as f:
        while chunk := f.read(65536):
            hasher.update(chunk)
    return hasher.hexdigest()

def get_content_fingerprint(pdf_path: str) -> str:
    try:
        doc = fitz.open(pdf_path)
        pages_to_check = min(3, len(doc))
        text = ""
        for i in range(pages_to_check):
            text += doc[i].get_text()
        doc.close()
        normalized = re.sub(r'[\s\W]+', ' ', text.lower()).strip()
        if len(normalized) < 50:
            return None
        return hashlib.sha256(normalized[:3000].encode("utf-8")).hexdigest()
    except Exception as e:
        print(f"  [Fingerprint] Could not extract content fingerprint: {e}")
        return None

def register_book_in_db(pdf_path: str, db: dict, file_hash: str,
                         content_fp: str, metadata: dict):
    """Registers a successfully processed book including metadata for future LLM dedup."""
    filename = os.path.basename(pdf_path)
    db[filename] = {
        "file_hash":           file_hash,
        "content_fingerprint": content_fp,
        "title_latin":         metadata.get("title_latin", ""),
        "title_persian":       metadata.get("title_persian", ""),
        "author_latin":        metadata.get("author_latin", ""),
        "author_persian":      metadata.get("author_persian", ""),
        "translator":          metadata.get("translator", ""),
        "publish_date":        metadata.get("publish_date", ""),
        "processed_at":        time.strftime("%Y-%m-%d %H:%M:%S"),
        "original_path":       pdf_path,
    }
    save_fingerprint_db(db)
    print(f"  [Dedup] Registered '{filename}' in fingerprint DB.")

# ==========================================
# DUPLICATE DETECTION
# ==========================================
def check_duplicate_via_llm(new_book: dict, existing_books: dict) -> tuple[bool, str]:
    if not existing_books:
        return False, "DB is empty"

    known_list = []
    for i, (filename, data) in enumerate(existing_books.items()):
        known_list.append(
            f"{i+1}. filename='{filename}' | "
            f"title_latin='{data.get('title_latin','')}' | "
            f"title_persian='{data.get('title_persian','')}' | "
            f"author_latin='{data.get('author_latin','')}' | "
            f"author_persian='{data.get('author_persian','')}' | "
            f"translator='{data.get('translator','')}' | "
            f"publish_date='{data.get('publish_date','')}'"
        )

    new_book_text = (
        f"title_latin='{new_book.get('title_latin','')}' | "
        f"title_persian='{new_book.get('title_persian','')}' | "
        f"author_latin='{new_book.get('author_latin','')}' | "
        f"author_persian='{new_book.get('author_persian','')}' | "
        f"translator='{new_book.get('translator','')}' | "
        f"publish_date='{new_book.get('publish_date','')}'"
    )

    prompt = f"""
You are a librarian expert in philosophy books, Persian translations, and academic publishing.

A new book has been scanned. Compare it against the list of already-processed books and determine
if the new book is a DUPLICATE of any existing entry.

RULES — counts as duplicate:
- Same book, author name written differently (e.g. "G.W.F. Hegel" vs "Georg Wilhelm Friedrich Hegel")
- Same book title in different language but same content
- Same book, slightly different edition year or publisher (reprint of same translation)
- Same translator + same original work = duplicate even if publisher differs

RULES — NOT a duplicate:
- Same author, different book title
- Same title but DIFFERENT translator (different translation = different book worth keeping)
- Same title, different language edition for different audience

NEW BOOK:
{new_book_text}

ALREADY PROCESSED BOOKS:
{chr(10).join(known_list)}

Respond ONLY with JSON:
{{
    "is_duplicate": true or false,
    "matched_entry_number": <number or null>,
    "matched_filename": "<filename or null>",
    "confidence": "high" or "medium" or "low",
    "reason": "<one sentence>"
}}
"""

    try:
        result = call_llm(prompt, context="duplicate check")
    except RuntimeError as e:
        print(f"  [Dedup-LLM] All LLMs failed: {e}. Assuming not duplicate.")
        return False, "LLM unavailable"

    is_dup     = result.get("is_duplicate", False)
    confidence = result.get("confidence", "low")
    reason     = result.get("reason", "")
    matched    = result.get("matched_filename", "")

    if is_dup and confidence == "low":
        print(f"  [Dedup-LLM] Low-confidence match ignored: {reason}")
        return False, f"low-confidence skipped: {reason}"

    if is_dup:
        return True, f"[{confidence}] matched '{matched}': {reason}"
    return False, reason


def is_duplicate_book(pdf_path: str, db: dict,
                       metadata: dict = None) -> tuple[bool, str]:
    """
    Layer 1 — exact filename in DB         (instant)
    Layer 2 — file hash                    (catches renames)
    Layer 3 — content fingerprint          (catches re-scans)
    Layer 4 — LLM semantic check           (catches typos, initials, language variants)
    Layer 5 — legacy .docx/.md on disk     (catches pre-DB processed books)
    """
    filename  = os.path.basename(pdf_path)
    base_name = os.path.splitext(filename)[0]

    # Layer 1
    if filename in db:
        return True, f"filename '{filename}' already in DB"

    # Layer 2
    print(f"  [Dedup] Computing file hash...")
    file_hash = get_file_hash(pdf_path)
    for stored_fn, stored_data in db.items():
        if stored_data.get("file_hash") == file_hash:
            return True, f"identical file already processed as '{stored_fn}'"

    # Layer 3
    print(f"  [Dedup] Computing content fingerprint...")
    content_fp = get_content_fingerprint(pdf_path)
    if content_fp:
        for stored_fn, stored_data in db.items():
            if stored_data.get("content_fingerprint") == content_fp:
                return True, f"same page content as '{stored_fn}'"

    # Layer 4 — only runs when metadata is available (post-OCR phase)
    if metadata and db:
        print(f"  [Dedup] Running LLM duplicate check against {len(db)} known books...")
        is_dup, reason = check_duplicate_via_llm(metadata, db)
        if is_dup:
            return True, f"LLM: {reason}"

    # Layer 5
    if os.path.exists(OUTPUT_DIR):
        for root, dirs, files_on_disk in os.walk(OUTPUT_DIR):
            for f in files_on_disk:
                f_base = os.path.splitext(f)[0].lower()
                f_ext  = os.path.splitext(f)[1].lower()
                if f_base == base_name.lower() and f_ext in ('.docx', '.md'):
                    return True, f"output file on disk: {os.path.join(root, f)}"

    return False, ""

# ==========================================
# HELPERS
# ==========================================
def escape_sql_string(val) -> str:
    if val is None:
        return "NULL"
    escaped = str(val).replace('\\', '\\\\').replace("'", "\\'")
    return f"'{escaped}'"

def retry_api_call(func, *args, max_retries=4, base_delay=10, **kwargs):
    current_delay = base_delay
    for attempt in range(max_retries):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            print(f"  [Warning] Attempt {attempt + 1} failed: {e}.")
            if attempt < max_retries - 1:
                print(f"  Retrying in {current_delay}s...")
                time.sleep(current_delay)
                current_delay *= 2
            else:
                raise e

def add_anchors_and_toc(markdown_text: str) -> str:
    """Preserves Persian/Arabic characters in anchors, deduplicates repeated headings."""
    print("Post-processing Markdown: Generating ToC...")
    lines = markdown_text.split('\n')
    toc = ["# Table of Contents\n"]
    new_lines = []
    anchor_counts = {}  # track duplicates

    for line in lines:
        stripped = line.strip()
        match = re.match(r'^(#{1,3})\s+(.*)', stripped)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()

            # Keep Persian/Arabic AND Latin alphanumerics, replace rest with -
            anchor_id = re.sub(r'[^\w\u0600-\u06FF\u0750-\u077F]', '-', title.lower())
            anchor_id = re.sub(r'-+', '-', anchor_id).strip('-') or "section"

            # Deduplicate anchors (same heading text appears multiple times)
            if anchor_id in anchor_counts:
                anchor_counts[anchor_id] += 1
                anchor_id = f"{anchor_id}-{anchor_counts[anchor_id]}"
            else:
                anchor_counts[anchor_id] = 0

            indent = "  " * (level - 1)
            toc.append(f"{indent}- [{title}](#{anchor_id})")
            new_lines.append(f"{match.group(1)} <a id=\"{anchor_id}\"></a> {title}")
        else:
            new_lines.append(line)

    return "\n".join(toc) + "\n\n---\n\n" + "\n".join(new_lines)

def save_markdown_as_docx(md_text: str, docx_path: str):
    doc = Document()
    lines = md_text.split('\n')
    in_table = False
    table_data = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            if re.match(r'^\|[\s:-|]+\|$', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_data:
                rows = len(table_data)
                cols = max(len(row) for row in table_data) if table_data else 0
                if cols > 0:
                    word_table = doc.add_table(rows=rows, cols=cols)
                    word_table.style = 'Table Grid'
                    for r_idx, row in enumerate(table_data):
                        for c_idx, val in enumerate(row):
                            if c_idx < len(word_table.rows[r_idx].cells):
                                word_table.rows[r_idx].cells[c_idx].text = val
                table_data = []
                in_table = False
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif not stripped:
            continue
        else:
            doc.add_paragraph(stripped)
    if in_table and table_data:
        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0
        if cols > 0:
            word_table = doc.add_table(rows=rows, cols=cols)
            word_table.style = 'Table Grid'
            for r_idx, row in enumerate(table_data):
                for c_idx, val in enumerate(row):
                    if c_idx < len(word_table.rows[r_idx].cells):
                        word_table.rows[r_idx].cells[c_idx].text = val
    doc.save(docx_path)

# ==========================================
# METADATA EXTRACTION
# ==========================================
METADATA_PROMPT_TEMPLATE = """
Analyze the filename and the text provided to determine the publishing details of this philosophy book.

CRITICAL INSTRUCTION FOR TRANSLATIONS: If the book is a translated text (e.g. German/Greek/French into Persian),
and the translator's name is not explicitly written in the front matter, use your broader academic and
historical knowledge to identify who translated this specific edition (for example, identifying famous Persian
philosophy translators like Javad Tabataba'i, Hamid Enayat, Seyyed Hossein Nasr, etc.). If you are reasonably
confident, output their name. If completely unknown, return null.

Filename: {filename}
Book Text Sample:
{text}

Your response must be a structured JSON object containing:
- title_latin
- title_persian
- author_latin
- author_persian
- translator
- publish_date
- publisher
- description (2-3 sentences)
- wp_tags (array of 5-10 relevant English keyword tags)
- wp_category (single best-fit category string)
"""

def extract_metadata(filename: str, front_matter_text: str) -> dict:
    print("Extracting metadata...")
    prompt = METADATA_PROMPT_TEMPLATE.format(
        filename=filename,
        text=front_matter_text[:4000]
    )
    try:
        return call_llm(prompt, context="metadata extraction")
    except RuntimeError as e:
        print(f"  [Metadata] All LLMs failed: {e}. Using filename defaults.")
        return {
            "title_latin":    filename,
            "title_persian":  None,
            "author_latin":   "Unknown",
            "author_persian": None,
            "translator":     None,
            "publish_date":   "Unknown",
            "publisher":      "Unknown",
            "description":    "No description available.",
            "wp_tags":        [],
            "wp_category":    "Uncategorized",
        }

# ==========================================
# SQL GENERATION
# ==========================================
def append_to_sql_dump_file(metadata: dict, language: str, md_text: str,
                             md_path: str, docx_path: str):
    print(f"Appending SQL to dump file...")

    # Store only: Persian/filename.md — server resolves the rest
    lang_folder   = {"fa": "Persian", "de": "German", "en": "English"}.get(language, "Other")
    md_filename   = os.path.basename(md_path)
    docx_filename = os.path.basename(docx_path)
    rel_md_path   = f"{lang_folder}/{md_filename}"
    rel_docx_path = f"{lang_folder}/{docx_filename}"

    raw_slug = metadata.get("title_latin") or metadata.get("title_persian") or "book"
    wp_slug  = re.sub(r'[^a-z0-9]+', '-', raw_slug.lower()).strip('-')

    def e(key): return escape_sql_string(metadata.get(key))

    sql = []

    sql.append(
        f"INSERT IGNORE INTO books "
        f"(title_latin, title_persian, author_latin, author_persian, "
        f"translator, publish_date, publisher, language, description, "
        f"md_file_path, docx_file_path, wp_slug)\n"
        f"VALUES ({e('title_latin')}, {e('title_persian')}, {e('author_latin')}, "
        f"{e('author_persian')}, {e('translator')}, {e('publish_date')}, "
        f"{e('publisher')}, {escape_sql_string(language)}, {e('description')}, "
        f"{escape_sql_string(rel_md_path)}, {escape_sql_string(rel_docx_path)}, "
        f"{escape_sql_string(wp_slug)});\n"
    )

    sql.append(
        f"SET @last_book_id = (SELECT id FROM books WHERE wp_slug = {escape_sql_string(wp_slug)});\n"
    )

    for section in re.split(r'\n(?=#{1,3}\s+)', md_text):
        title_match   = re.match(r'^(#{1,3})\s+(.*)', section.strip())
        section_title = title_match.group(2).strip() if title_match else "Introduction / Prologue"
        sql.append(
            f"INSERT IGNORE INTO book_contents (book_id, section_title, text_content)\n"
            f"VALUES (@last_book_id, {escape_sql_string(section_title)}, "
            f"{escape_sql_string(section.strip())});\n"
        )

    wp_tags_json = escape_sql_string(json.dumps(metadata.get("wp_tags", []), ensure_ascii=False))
    sql.append(
        f"INSERT IGNORE INTO wp_sync "
        f"(book_id, wp_post_title, wp_slug, wp_tags_json, wp_category, wp_status, synced)\n"
        f"VALUES (@last_book_id, {e('title_latin')}, {escape_sql_string(wp_slug)}, "
        f"{wp_tags_json}, {e('wp_category')}, 'draft', 0);\n"
    )

    with open(SQL_DUMP_FILE_PATH, "a", encoding="utf-8") as f:
        f.write("\n" + "".join(sql) + "\n")


# ==========================================
# PDF PROCESSING
# ==========================================
def split_pdf_into_chunks(pdf_path: str, chunk_size_pages: int) -> list:
    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    chunk_paths = []
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    temp_dir  = os.path.dirname(pdf_path)
    print(f"Splitting {base_name} ({total_pages} pages) into {chunk_size_pages}-page chunks...")
    for start_page in range(0, total_pages, chunk_size_pages):
        end_page   = min(start_page + chunk_size_pages - 1, total_pages - 1)
        chunk_doc  = fitz.open()
        chunk_doc.insert_pdf(doc, from_page=start_page, to_page=end_page)
        chunk_path = os.path.join(temp_dir, f"temp_{base_name}_part_{start_page//chunk_size_pages+1}.pdf")
        chunk_doc.save(chunk_path)
        chunk_doc.close()
        chunk_paths.append(chunk_path)
    doc.close()
    return chunk_paths

def process_single_pdf_to_markdown_with_retry(pdf_path: str) -> str:
    def upload_action():
        print(f"Uploading '{os.path.basename(pdf_path)}' to Mistral...")
        with open(pdf_path, "rb") as f:
            return client.files.upload(
                file={"file_name": os.path.basename(pdf_path), "content": f.read()},
                purpose="ocr"
            )
    uploaded_file = retry_api_call(upload_action)
    file_id = uploaded_file.id
    print(f"Upload OK (ID: {file_id}).")
    def ocr_action():
        print("Running Mistral OCR...")
        return client.ocr.process(
            model="mistral-ocr-latest",
            document={"type": "file", "file_id": file_id}
        )
    try:
        ocr_response = retry_api_call(ocr_action)
        return "\n\n".join([page.markdown for page in ocr_response.pages])
    finally:
        try:
            client.files.delete(file_id=file_id)
            print(f"Deleted cloud file {file_id}.")
        except Exception as e:
            print(f"Cloud cleanup warning: {e}")

def detect_book_language(pdf_path: str) -> str:
    print(f"\n--- Language detection: {os.path.basename(pdf_path)} ---")
    temp_mid_path = os.path.join(os.path.dirname(pdf_path), f"temp_lang_check_{os.path.basename(pdf_path)}")
    doc = fitz.open(pdf_path)
    total_pages  = len(doc)
    mid_page_idx = max(0, min(total_pages // 2, total_pages - 1))
    print(f"Using midpoint page {mid_page_idx + 1} of {total_pages}...")
    p1_doc = fitz.open()
    p1_doc.insert_pdf(doc, from_page=mid_page_idx, to_page=mid_page_idx)
    p1_doc.save(temp_mid_path)
    p1_doc.close()
    doc.close()
    try:
        p1_markdown = process_single_pdf_to_markdown_with_retry(temp_mid_path)
        def chat_action():
            return client.chat.complete(
                model="mistral-small-latest",
                messages=[{"role": "user", "content": f"Respond ONLY with the ISO 639-1 code of the language:\n\n{p1_markdown}"}]
            )
        chat_response  = retry_api_call(chat_action)
        detected_lang  = re.sub(r'[^a-z]', '', chat_response.choices[0].message.content.strip().lower())[:2]
        print(f"--> Language: {detected_lang.upper()}")
        return detected_lang
    except Exception as e:
        print(f"Language detection failed: {e}. Defaulting to 'en'.")
        return "en"
    finally:
        if os.path.exists(temp_mid_path):
            os.remove(temp_mid_path)

def run_ocr_on_book(pdf_path: str, final_md_path: str, final_docx_path: str,
                     lang_code: str) -> tuple[dict, str]:
    """
    Runs OCR only — writes .md and .docx, extracts metadata.
    Does NOT write SQL (that's main()'s job after dedup passes).
    Returns (metadata, combined_markdown).
    """
    file_size_mb = os.path.getsize(pdf_path) / (1024 * 1024)
    doc_meta     = fitz.open(pdf_path)
    total_pages  = len(doc_meta)
    doc_meta.close()
    print(f"File: {file_size_mb:.2f} MB, {total_pages} pages")

    temp_chunks = (
        split_pdf_into_chunks(pdf_path, PAGES_PER_CHUNK)
        if file_size_mb > MAX_FILE_SIZE_MB or total_pages > MAX_PAGES_PER_MISTRAL
        else [pdf_path]
    )

    full_book_markdown = []
    try:
        for idx, chunk in enumerate(temp_chunks):
            print(f"Processing chunk {idx+1}/{len(temp_chunks)}")
            chunk_md = process_single_pdf_to_markdown_with_retry(chunk)
            full_book_markdown.append(chunk_md)
            if idx < len(temp_chunks) - 1:
                print(f"Sleeping {SLEEP_BETWEEN_CHUNKS}s between chunks...")
                time.sleep(SLEEP_BETWEEN_CHUNKS)

        combined_markdown = "\n\n".join(full_book_markdown)
        combined_markdown = add_anchors_and_toc(combined_markdown)

        metadata = extract_metadata(os.path.basename(pdf_path), combined_markdown[:6000])
        print(
            f"Metadata — Title: {metadata.get('title_latin')} | "
            f"Author: {metadata.get('author_latin')} | "
            f"Translator: {metadata.get('translator')} | "
            f"WP Category: {metadata.get('wp_category')}"
        )

        # Write files
        with open(final_md_path, "w", encoding="utf-8") as md_file:
            md_file.write(combined_markdown)
        print(f"  ✓ Saved MD:   {final_md_path}")

        save_markdown_as_docx(combined_markdown, final_docx_path)
        print(f"  ✓ Saved DOCX: {final_docx_path}")

        # Return both — SQL is written by main() only after dedup passes
        return metadata, combined_markdown

    finally:
        if len(temp_chunks) > 1:
            print("Cleaning up temp chunk files...")
            for chunk in temp_chunks:
                if chunk != pdf_path and os.path.exists(chunk):
                    os.remove(chunk)


# ==========================================
# MAIN
# ==========================================
def main():
    if not os.path.exists(INPUT_BOOKS_DIR):
        print(f"Input directory not found: {INPUT_BOOKS_DIR}")
        return

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    fingerprint_db = load_fingerprint_db()
    print(f"[Dedup] Loaded fingerprint DB — {len(fingerprint_db)} known books.")

    files = [
        os.path.join(INPUT_BOOKS_DIR, f)
        for f in os.listdir(INPUT_BOOKS_DIR)
        if f.lower().endswith('.pdf')
    ]
    print(f"Found {len(files)} PDF(s).")

    for idx, pdf_file in enumerate(files):
        base_name = os.path.splitext(os.path.basename(pdf_file))[0]
        print(f"\n{'='*60}")
        print(f"[{idx+1}/{len(files)}] {base_name}")
        print(f"{'='*60}")

        # === PHASE 1: fast checks BEFORE touching any API ===
        is_dup, reason = is_duplicate_book(pdf_file, fingerprint_db, metadata=None)
        if is_dup:
            print(f"[Skip — Pre-OCR] {reason}")
            continue

        # Pre-compute hashes while file is cold (before upload)
        file_hash  = get_file_hash(pdf_file)
        content_fp = get_content_fingerprint(pdf_file)

        try:
            lang_code       = detect_book_language(pdf_file)
            lang_subfolder  = {"fa": "Persian", "de": "German", "en": "English"}.get(lang_code, "Other")
            lang_output_dir = os.path.join(OUTPUT_DIR, lang_subfolder)
            os.makedirs(lang_output_dir, exist_ok=True)

            final_md_path   = os.path.join(lang_output_dir, f"{base_name}.md")
            final_docx_path = os.path.join(lang_output_dir, f"{base_name}.docx")

            # OCR — writes .md and .docx, returns (metadata, markdown)
            metadata, combined_markdown = run_ocr_on_book(
                pdf_file, final_md_path, final_docx_path, lang_code
            )

            # === PHASE 2: LLM semantic dedup with metadata ===
            is_dup, reason = is_duplicate_book(pdf_file, fingerprint_db, metadata=metadata)
            if is_dup:
                print(f"[Skip — Post-OCR Dedup] {reason}")
                print(f"  Removing redundant output files...")
                for path in [final_md_path, final_docx_path]:
                    if os.path.exists(path):
                        os.remove(path)
                        print(f"  Deleted: {path}")
                continue

            # === ALL CLEAR — write SQL, register in DB ===
            # SQL only written here, after both dedup phases pass
            append_to_sql_dump_file(
                metadata, lang_code, combined_markdown, final_md_path, final_docx_path
            )

            register_book_in_db(pdf_file, fingerprint_db, file_hash, content_fp, metadata)

            print(f"\n✓ Book complete:")
            print(f"  MD:   {final_md_path}")
            print(f"  DOCX: {final_docx_path}")

            if idx < len(files) - 1:
                print(f"Sleeping {SLEEP_BETWEEN_BOOKS}s before next book...")
                time.sleep(SLEEP_BETWEEN_BOOKS)

        except Exception as e:
            print(f"CRITICAL ERROR on '{base_name}': {e}. Will retry next run.")
            # Files may or may not exist — do NOT delete them here.
            # If they exist but DB registration failed, next run's Layer 5
            # (disk check) will catch them and skip. The book won't be
            # reprocessed wastefully.

if __name__ == "__main__":
    main()
